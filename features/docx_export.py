"""
features/docx_export.py
DOCX Export
Builds a formatted Word document from a script dict and returns a Flask send_file response.
"""

import logging
from io import BytesIO
from flask import send_file, jsonify
from docx import Document

logger = logging.getLogger(__name__)


class DocxExporter:
    """
    Converts a script data dict into a downloadable DOCX file.
    Resolves script data from in-memory cache first, then DB fallback.
    """

    def __init__(self, script_store, scripts_cache: dict):
        """
        Args:
            script_store:  features.script_store.ScriptStore instance
            scripts_cache: module-level dict {script_id: script_data}
                           populated at generation time (in-memory, session only)
        """
        self.script_store  = script_store
        self.scripts_cache = scripts_cache

    def export(self, script_id: str):
        """
        Look up script by ID and return a Flask send_file response.

        Lookup order:
            1. In-memory cache (scripts generated this session)
            2. Database via ScriptStore

        Returns:
            Flask response — either a DOCX file download or a JSON error.
        """
        # 1. Cache
        script_data = self.scripts_cache.get(script_id)

        # 2. DB fallback
        if not script_data:
            script_data = self.script_store.get_script_by_id(script_id)

        if not script_data:
            return jsonify({'error': 'Script not found'}), 404

        docx_io = self._build(script_data)
        if not docx_io:
            return jsonify({'error': 'DOCX generation failed'}), 500

        return send_file(
            docx_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'script_{script_id}.docx'
        )

    def _build(self, script_data: dict) -> BytesIO | None:
        """
        Build the DOCX in memory and return a seeked BytesIO buffer.

        Sections:
            - Title (heading 1)
            - Metadata: subject, topic, word count, generated_at
            - Hook
            - Script content
            - Callaway direction (if present)
            - Callaway lens (if present)
            - Thumbnail brief (if present)
        """
        try:
            doc = Document()

            doc.add_heading(script_data.get('title', 'Untitled'), level=1)

            doc.add_paragraph(f"Subject: {script_data.get('subject', 'N/A')}")
            doc.add_paragraph(f"Topic: {script_data.get('topic', 'N/A')}")
            doc.add_paragraph(f"Word Count: {script_data.get('word_count', 0)}")
            doc.add_paragraph(f"Generated: {script_data.get('generated_at', 'N/A')}")

            doc.add_heading("Hook", level=2)
            doc.add_paragraph(script_data.get('hook', ''))

            doc.add_heading("Script", level=2)
            doc.add_paragraph(script_data.get('script_content', ''))

            if script_data.get('callaway_direction'):
                doc.add_heading("Direction", level=2)
                doc.add_paragraph(str(script_data['callaway_direction']))

            if script_data.get('callaway_lens'):
                doc.add_heading("Lens", level=2)
                doc.add_paragraph(str(script_data['callaway_lens']))

            if script_data.get('thumbnail_prompt'):
                doc.add_heading("Thumbnail Brief", level=2)
                doc.add_paragraph(script_data['thumbnail_prompt'])

            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf

        except Exception as e:
            logger.error(f"DOCX build error: {e}")
            return None

    def export_long_form(self, subject: str, topic: str, script_content: str, outline: dict = None, metadata: dict = None) -> bytes:
        """
        Export a long-form video script as DOCX.
        
        Args:
            subject: Subject name
            topic: Topic name
            script_content: Full script text
            outline: Optional outline dict with phases and sections
            metadata: Optional metadata dict with word_count, duration_minutes, etc
        
        Returns:
            BytesIO buffer ready for send_file()
        """
        try:
            doc = Document()
            
            metadata = metadata or {}
            
            # Title
            doc.add_heading(f"{subject}: {topic}", level=1)
            
            # Metadata
            doc.add_paragraph(f"Type: Long-Form Video", style='Normal')
            doc.add_paragraph(f"Duration: {metadata.get('duration_minutes', 'N/A')} minutes", style='Normal')
            doc.add_paragraph(f"Word Count: {metadata.get('word_count', 0):,}", style='Normal')
            if metadata.get('validation'):
                doc.add_paragraph(f"Curriculum Coverage: {metadata['validation'].get('coverage_percentage', 0)}%", style='Normal')
            
            # Outline section
            if outline:
                doc.add_heading("Outline", level=2)
                doc.add_paragraph(f"Thesis: {outline.get('thesis_question', 'N/A')}")
                doc.add_paragraph(f"Framework: {outline.get('metaphor', 'N/A')}")
                
                if 'sections' in outline:
                    doc.add_heading("Sections", level=3)
                    for i, section in enumerate(outline['sections'], 1):
                        doc.add_paragraph(
                            f"{i}. {section.get('name', f'Section {i}')} (~{section.get('target_words', 0)} words)",
                            style='List Number'
                        )
            
            # Full script
            doc.add_heading("Full Script", level=2)
            doc.add_paragraph(script_content)
            
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.getvalue()
        
        except Exception as e:
            logger.error(f"Long-form DOCX export error: {e}")
            return None

    def export_essay(self, title: str, content: str, metadata: dict = None) -> bytes:
        """
        Export an essay as DOCX.
        
        Args:
            title: Essay title
            content: Full essay content (markdown formatted with # and ## headers)
            metadata: Optional metadata dict with essay_type, tone, word_count, etc
        
        Returns:
            BytesIO buffer ready for send_file()
        """
        try:
            doc = Document()
            
            metadata = metadata or {}
            
            # Title
            doc.add_heading(title, level=1)
            
            # Metadata
            doc.add_paragraph(f"Type: {metadata.get('essay_type', 'essay').title()}", style='Normal')
            doc.add_paragraph(f"Tone: {metadata.get('tone', 'N/A').title()}", style='Normal')
            doc.add_paragraph(f"Word Count: {metadata.get('word_count', 0):,}", style='Normal')
            doc.add_paragraph(f"Target Audience: {metadata.get('target_audience', 'N/A')}", style='Normal')
            
            # Parse and add content
            lines = content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    # Top-level heading (skip, already have title)
                    continue
                elif line.startswith('## '):
                    # Second-level heading
                    heading_text = line.replace('## ', '').strip()
                    doc.add_heading(heading_text, level=2)
                elif line.startswith('### '):
                    # Third-level heading
                    heading_text = line.replace('### ', '').strip()
                    doc.add_heading(heading_text, level=3)
                elif line.strip():
                    # Regular paragraph
                    doc.add_paragraph(line.strip())
                else:
                    # Empty line - add spacing
                    doc.add_paragraph()
            
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.getvalue()
        
        except Exception as e:
            logger.error(f"Essay DOCX export error: {e}")
            return None
